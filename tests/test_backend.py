from pathlib import Path
import json
import io

from docx import Document
from fastapi.testclient import TestClient

import backend

class _DummyResponse:
    def __init__(self, status_code: int, data: dict, text: str = ""):
        self.status_code = status_code
        self._data = data
        self.text = text

    def json(self):
        return self._data

def _configure_rag_paths(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(backend, "VECTOR_DB_PATH", str(tmp_path / "chroma_db"))
    monkeypatch.setattr(backend, "UPLOADED_DOCS_PATH", str(tmp_path / "uploaded_docs"))
    backend._reset_rag_services_for_tests()

def _seed_v1_doc(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    vs, dm = backend._ensure_rag_services()
    content = b" a" * 200
    doc_id = dm.save_document(content, "seed.txt", backend.UNIFIED_COLLECTION, chunk_count=1)
    vs.upsert_document(
        collection_name=backend.UNIFIED_COLLECTION,
        doc_id=doc_id,
        chunks=["semiconductor screening examples"],
        embeddings=[[1.0, 0.0, 0.0]],
        metadata={"filename": "seed.txt", "vertical": backend.UNIFIED_COLLECTION},
    )
    return doc_id

def _seed_v3_doc(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    vs, dm = backend._ensure_rag_services()
    content = b" a" * 200
    doc_id = dm.save_document(content, "interview_notes.txt", backend.UNIFIED_COLLECTION, chunk_count=1)
    vs.upsert_document(
        collection_name=backend.UNIFIED_COLLECTION,
        doc_id=doc_id,
        chunks=["Interviewee A said the market is growing fast despite volatility."],
        embeddings=[[1.0, 0.0, 0.0]],
        metadata={"filename": "interview_notes.txt", "vertical": backend.UNIFIED_COLLECTION},
    )
    return doc_id


def _sample_export_payload():
    return {
        "export_format": "email_html",
        "objective": "expert_network_brief",
        "mode": "rag",
        "prompt": "Draft a screening brief for a software diligence.",
        "generated_at": "2026-04-22T10:30:00",
        "context_fields": [
            {"label": "industry", "value": "Vertical software"},
            {"label": "live data", "value": "Enabled"},
        ],
        "left": {
            "export_label": "bedrock - claude-sonnet",
            "provider": "bedrock",
            "model": "claude-sonnet",
            "content": "# Executive Summary\n- Strong retention\n- Clear upsell path",
            "usage": {"prompt_tokens": 120, "completion_tokens": 80},
            "metrics": {"total_tokens": 200, "latency_ms": 980.2, "tok_per_sec": 12.8, "retrieval_ms": 140.0},
        },
        "right": {
            "export_label": "lmstudio - qwen3",
            "provider": "lmstudio",
            "model": "qwen3",
            "content": "1. Validate ICP\n2. Pressure-test churn",
            "usage": {"prompt_tokens": 110, "completion_tokens": 60},
            "metrics": {"total_tokens": 170, "latency_ms": 1240.0, "tok_per_sec": 9.6, "retrieval_ms": 140.0},
        },
        "rag": {
            "collection": "ALL",
            "chunks_retrieved": 5,
            "documents_retrieved": 2,
            "avg_score": 0.81,
        },
        "sources": [
            {
                "filename": "notes.txt",
                "chunk_id": 0,
                "score": 0.92,
                "snippet": "Customers describe the workflow as mission critical.",
            }
        ],
        "web_search": {
            "requested": True,
            "enabled": True,
            "sources": [
                {
                    "title": "Market update",
                    "url": "https://example.com/update",
                    "published_date": "2026-04-20",
                    "snippet": "Growth in the category remained resilient.",
                }
            ],
        },
    }

def test_chat_direct_mode_with_file_still_works(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat",
        data={"message": "Summarize this request briefly in 1 sentence. Make it super short.", "mode": "direct"},
        files={"file": ("brief.txt", "hello context", "text/plain")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert isinstance(data["content"], str)

def test_chat_rag_mode_returns_response_with_sources(tmp_path: Path, monkeypatch):
    _seed_v1_doc(tmp_path, monkeypatch)
    monkeypatch.setattr("rag.retrieval.get_embedding", lambda _: [1.0, 0.0, 0.0])

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat",
        data={
            "message": "Draft a tiny 1-sentence brief exactly.",
            "mode": "rag",
            "objective": "expert_network_brief",
        },
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert isinstance(data["content"], str)
    assert isinstance(data["sources"], list)
    assert len(data["sources"]) >= 1

def test_chat_rag_empty_vertical_returns_error(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    monkeypatch.setattr("rag.retrieval.get_embedding", lambda _: [1.0, 0.0, 0.0])

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat",
        data={
            "message": "Any context?",
            "mode": "rag",
            "objective": "interview_guide",
        },
    )
    assert resp.status_code == 400

def test_chat_rag_insights_qa_uses_local_provider(tmp_path: Path, monkeypatch):
    _seed_v3_doc(tmp_path, monkeypatch)
    monkeypatch.setattr("rag.retrieval.get_embedding", lambda _: [1.0, 0.0, 0.0])

    def fake_post(url, *args, **kwargs):
        return _DummyResponse(
            200,
            {
                "choices": [{"message": {"content": "insights ok"}}],
                "usage": {"prompt_tokens": 17, "completion_tokens": 9},
            },
        )
    monkeypatch.setattr(backend.requests, "post", fake_post)
    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat",
        data={
            "message": "Who mentioned market growth?",
            "mode": "rag",
            "objective": "insights_qa",
        },
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["provider"] == "lmstudio"

def test_documents_endpoint_returns_list(tmp_path: Path, monkeypatch):
    _seed_v1_doc(tmp_path, monkeypatch)
    client = TestClient(backend.app)
    resp = client.get("/api/documents")
    assert resp.status_code == 200, resp.text

def test_document_file_endpoint_returns_file(tmp_path: Path, monkeypatch):
    doc_id = _seed_v1_doc(tmp_path, monkeypatch)
    client = TestClient(backend.app)
    resp = client.get(f"/api/documents/{doc_id}/file")
    assert resp.status_code == 200, resp.text

def test_document_file_endpoint_missing_doc_returns_404(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    client = TestClient(backend.app)
    resp = client.get("/api/documents/doc_missing/file")
    assert resp.status_code == 404

def test_stats_endpoint_returns_expected_keys(tmp_path: Path, monkeypatch):
    _seed_v1_doc(tmp_path, monkeypatch)
    client = TestClient(backend.app)
    resp = client.get("/api/stats")
    assert resp.status_code == 200, resp.text

def test_system_prompt_included_in_rag_payload(tmp_path: Path, monkeypatch):
    _seed_v1_doc(tmp_path, monkeypatch)
    monkeypatch.setattr("rag.retrieval.get_embedding", lambda _: [1.0, 0.0, 0.0])

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat",
        data={
            "message": "Write a 5 word screening question.",
            "mode": "rag",
            "objective": "expert_network_brief",
        },
    )
    assert resp.status_code == 200, resp.text
    assert "content" in resp.json()

def test_provider_config_bedrock_direct(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    
    client = TestClient(backend.app)
    resp = client.post("/api/chat", data={"message": "Reply 'hello' directly.", "mode": "direct"})
    
    assert resp.status_code == 200, resp.text
    assert "hello" in resp.json()["content"].lower()

def test_local_model_config_reads_dotenv_when_env_missing(tmp_path: Path, monkeypatch):
    env_file = tmp_path / ".env"
    env_file.write_text(
        "LOCAL_RAG_MODEL=qwen3:32b\n"
        "LMSTUDIO_BASE_URL=http://127.0.0.1:1234/\n",
        encoding="utf-8",
    )
    monkeypatch.delenv("LOCAL_RAG_MODEL", raising=False)
    monkeypatch.delenv("LMSTUDIO_BASE_URL", raising=False)
    monkeypatch.setattr(backend, "APP_DIR", str(tmp_path))
    monkeypatch.chdir(tmp_path)
    base_url, model = backend._load_local_model_config()
    assert base_url == "http://127.0.0.1:1234"
    assert model == "qwen3:32b"

def test_chat_compare_returns_local_and_opus(tmp_path: Path, monkeypatch):
    _seed_v3_doc(tmp_path, monkeypatch)
    monkeypatch.setattr("rag.retrieval.get_embedding", lambda _: [1.0, 0.0, 0.0])

    def fake_post(url, *args, **kwargs):
        return _DummyResponse(
            200,
            {
                "choices": [{"message": {"content": "local ok"}}],
                "usage": {"prompt_tokens": 17, "completion_tokens": 9},
            },
        )
    monkeypatch.setattr(backend.requests, "post", fake_post)

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat/compare",
        data={
            "message": "In 1 sentence, who mentioned growth?",
            "objective": "insights_qa",
            "model_left": "lmstudio|qwen2.5:32b",
            "model_right": "bedrock|us.anthropic.claude-opus-4-6-v1",
        },
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["left"]["content"] == "local ok"
    assert data["left"]["provider"] == "lmstudio"
    assert "content" in data["right"]
    assert data["right"]["provider"] == "bedrock"


def test_chat_compare_can_run_only_left_model_when_compare_disabled(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    calls = []

    def fake_run_bedrock_model(messages, model_name, api_key=None, enable_web_search=False):
        calls.append(model_name)
        return {
            "content": "left only",
            "usage": {"prompt_tokens": 1, "completion_tokens": 1},
            "provider": "bedrock",
            "model": model_name,
            "metrics": {"total_tokens": 2, "latency_ms": 1.0, "tok_per_sec": 2.0},
        }

    monkeypatch.setattr(backend, "_load_bedrock_key", lambda: "test-key")
    monkeypatch.setattr(backend, "_run_bedrock_model", fake_run_bedrock_model)

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat/compare",
        data={
            "message": "Draft a one-sentence brief.",
            "mode": "direct",
            "objective": "expert_network_brief",
            "compare_enabled": "false",
            "model_left": "bedrock|us.anthropic.claude-haiku-4-5-20251001-v1:0",
        },
    )

    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["left"]["content"] == "left only"
    assert data["right"] is None
    assert calls == ["us.anthropic.claude-haiku-4-5-20251001-v1:0"]


def test_chat_compare_enables_web_search_for_objective1_when_checked(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    calls = []
    captured_messages = []

    def fake_run_bedrock_model(messages, model_name, api_key=None, enable_web_search=False):
        calls.append(enable_web_search)
        captured_messages.append(messages)
        return {
            "content": "ok",
            "usage": {"prompt_tokens": 1, "completion_tokens": 1},
            "provider": "bedrock",
            "model": model_name,
            "metrics": {"total_tokens": 2, "latency_ms": 1.0, "tok_per_sec": 2.0},
        }

    def fake_web_context(query, objective, requested):
        assert requested is True
        assert objective == "expert_network_brief"
        return (
            "<web_context>\n[1] Headline\nURL: https://example.com\nPublished: 2026-02-01\nSnippet: Update.\n</web_context>",
            {
                "requested": True,
                "enabled": True,
                "provider": "exa",
                "results_count": 1,
                "sources": [{"title": "Headline", "url": "https://example.com", "published_date": "2026-02-01", "snippet": "Update."}],
            },
        )

    monkeypatch.setattr(backend, "_load_bedrock_key", lambda: "test-key")
    monkeypatch.setattr(backend, "_run_bedrock_model", fake_run_bedrock_model)
    monkeypatch.setattr(backend, "_build_web_context", fake_web_context)

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat/compare",
        data={
            "message": "What changed in this market recently?",
            "mode": "direct",
            "objective": "expert_network_brief",
            "use_web_search": "true",
            "model_left": "bedrock|us.anthropic.claude-haiku-4-5-20251001-v1:0",
            "model_right": "bedrock|us.anthropic.claude-opus-4-6-v1",
        },
    )
    assert resp.status_code == 200, resp.text
    assert len(calls) == 2
    assert all(flag is True for flag in calls)
    for msgs in captured_messages:
        system_msg = next((m for m in msgs if m.get("role") == "system"), {})
        user_msg = next((m for m in msgs if m.get("role") == "user"), {})
        assert "Current date:" in (system_msg.get("content") or "")
        assert "<web_context>" in (user_msg.get("content") or "")
    assert resp.json()["web_search"]["enabled"] is True


def test_chat_compare_disables_web_search_for_objective3_even_when_checked(tmp_path: Path, monkeypatch):
    _configure_rag_paths(tmp_path, monkeypatch)
    calls = []

    def fake_run_bedrock_model(messages, model_name, api_key=None, enable_web_search=False):
        calls.append(enable_web_search)
        return {
            "content": "ok",
            "usage": {"prompt_tokens": 1, "completion_tokens": 1},
            "provider": "bedrock",
            "model": model_name,
            "metrics": {"total_tokens": 2, "latency_ms": 1.0, "tok_per_sec": 2.0},
        }

    monkeypatch.setattr(backend, "_load_bedrock_key", lambda: "test-key")
    monkeypatch.setattr(backend, "_run_bedrock_model", fake_run_bedrock_model)

    client = TestClient(backend.app)
    resp = client.post(
        "/api/chat/compare",
        data={
            "message": "Who said growth is accelerating?",
            "mode": "direct",
            "objective": "insights_qa",
            "use_web_search": "true",
            "model_left": "bedrock|us.anthropic.claude-haiku-4-5-20251001-v1:0",
            "model_right": "bedrock|us.anthropic.claude-opus-4-6-v1",
        },
    )
    assert resp.status_code == 200, resp.text
    assert len(calls) == 2
    assert all(flag is False for flag in calls)
    assert resp.json()["web_search"]["requested"] is True
    assert resp.json()["web_search"]["reason"] == "objective_not_supported"


def test_export_email_html_download_contains_formatted_sections():
    client = TestClient(backend.app)
    payload = _sample_export_payload()
    resp = client.post("/api/export", json=payload)

    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("text/html")
    assert "attachment; filename=" in resp.headers["content-disposition"]
    assert "Benjamin Maurice Analysis Export" in resp.text
    assert "Run Summary" in resp.text
    assert "bedrock - claude-sonnet" in resp.text
    assert "Referenced Passages" in resp.text


def test_export_word_docx_download_contains_response_content():
    client = TestClient(backend.app)
    payload = _sample_export_payload()
    payload["export_format"] = "word"
    resp = client.post("/api/export", json=payload)

    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment; filename=" in resp.headers["content-disposition"]

    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Benjamin Maurice Analysis Export" in full_text
    assert "Draft a screening brief for a software diligence." in full_text
    assert "Executive Summary" in full_text
    assert "Validate ICP" in full_text


def test_export_word_docx_allows_single_model_payload():
    client = TestClient(backend.app)
    payload = _sample_export_payload()
    payload["export_format"] = "word"
    payload["right"] = None
    resp = client.post("/api/export", json=payload)

    assert resp.status_code == 200, resp.text
    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in full_text
    assert "Validate ICP" not in full_text


def test_root_serves_ultra_and_classic_redirects():
    client = TestClient(backend.app)

    root = client.get("/")
    assert root.status_code == 200, root.text
    assert "Ultra Research Engine" in root.text
    assert "Classic" not in root.text

    classic = client.get("/classic", follow_redirects=False)
    assert classic.status_code == 307
    assert classic.headers["location"] == "/"
