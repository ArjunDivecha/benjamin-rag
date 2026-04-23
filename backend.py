"""
=============================================================================
SCRIPT NAME: backend.py
=============================================================================

INPUT FILES: None (receives file uploads via API)

OUTPUT FILES: None (returns JSON responses)

VERSION: 1.0
LAST UPDATED: 2026-02-11

DESCRIPTION:
FastAPI backend that proxies requests to OpenRouter Claude Sonnet 4.5
(Bedrock + ZDR). Supports direct file-context chat and RAG chat mode.

DEPENDENCIES:
- fastapi, uvicorn, python-multipart, requests

USAGE:
uvicorn backend:app --reload --port 8000

Then open http://localhost:8000
=============================================================================
"""

import io
import json
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Optional

import boto3
import requests
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles

from preprocess import _resolve_files, sync_collection

# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
APP_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(APP_DIR, "static")
AWS_REGION = "us-west-2"
MODEL = "us.anthropic.claude-opus-4-6-v1"
OPUS_COMPARE_MODEL = "us.anthropic.claude-opus-4-6-v1"
LMSTUDIO_BASE_URL = "http://localhost:1234"
LOCAL_RAG_MODEL = "qwen3:32b"
MAX_FILE_SIZE = 2 * 1024 * 1024  # 2MB
TEXT_ENCODINGS = ["utf-8", "latin-1", "cp1252"]
VECTOR_DB_PATH = os.path.join(APP_DIR, "chroma_db")
UPLOADED_DOCS_PATH = os.path.join(APP_DIR, "uploaded_docs")
UNIFIED_COLLECTION = "ALL"
OBJECTIVE_TO_COLLECTION = {
    "expert_network_brief": UNIFIED_COLLECTION,
    "interview_guide": UNIFIED_COLLECTION,
    "insights_qa": UNIFIED_COLLECTION,
}
OBJECTIVE_TO_PROVIDER = {
    "expert_network_brief": "bedrock",
    "interview_guide": "bedrock",
    "insights_qa": "lmstudio",
}
OBJECTIVE_LABELS = {
    "expert_network_brief": "Expert Brief",
    "interview_guide": "Interview Guide",
    "insights_qa": "Insights Engine",
}
MODE_LABELS = {
    "direct": "Attached File",
    "rag": "RAG Archive",
}
WEB_SEARCH_OBJECTIVES = {"expert_network_brief", "interview_guide"}
EXA_SEARCH_URL = "https://api.exa.ai/search"
EXA_SEARCH_RESULT_LIMIT = 5
EXA_SNIPPET_CHARS = 320
EXA_SEARCH_TIMEOUT_S = 15
INLINE_TOKEN_PATTERN = re.compile(
    r"(\*\*.+?\*\*|__.+?__|`.+?`|\*[^*\n][^*\n]*\*|_[^_\n][^_\n]*_)"
)
HEADING_PATTERN = re.compile(r"^\s*(#{1,6})\s+(.+?)\s*$")
UL_PATTERN = re.compile(r"^\s*[-*]\s+(.+?)\s*$")
OL_PATTERN = re.compile(r"^\s*\d+[.)]\s+(.+?)\s*$")

_vector_store = None
_doc_manager = None
_rag_import_error = None

app = FastAPI(title="Sonnet + File Context", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# -----------------------------------------------------------------------------
# Credentials
# -----------------------------------------------------------------------------
def _load_env_value_from_file(path: Path, key: str) -> Optional[str]:
    """Extract one key from a dotenv-style file."""
    if not path.exists():
        return None

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        if line.startswith("export "):
            line = line[len("export ") :].strip()
        if not line.startswith(f"{key}="):
            continue
        value = line.split("=", 1)[1].strip().strip('"\'')
        if value:
            return value
    return None


def _load_env_value(key: str) -> Optional[str]:
    """Load a value from environment, then local dotenv files."""
    value = os.environ.get(key, "").strip()
    if value:
        return value

    env_candidates = [
        Path(APP_DIR) / ".env",
        Path.cwd() / ".env",
        Path(APP_DIR) / ".env.txt",
        Path.cwd() / ".env.txt",
    ]
    for env_file in env_candidates:
        value = _load_env_value_from_file(env_file, key)
        if value:
            return value
    return None


def _load_bedrock_key() -> str:
    """Load Bedrock API key from environment or local dotenv files."""
    key = _load_env_value("BEDROCK_API_KEY")
    if key:
        return key

    raise RuntimeError(
        "BEDROCK_API_KEY not found. Create a local .env file with "
        "BEDROCK_API_KEY=your_key"
    )


def _load_exa_key() -> Optional[str]:
    """Load Exa API key from environment or local dotenv files."""
    return _load_env_value("EXA_API_KEY")


# -----------------------------------------------------------------------------
# File reading
# -----------------------------------------------------------------------------
def _read_docx(content: bytes) -> str:
    """Extract text from .docx file."""
    from io import BytesIO
    from docx import Document
    doc = Document(BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_file_content(content: bytes, filename: str) -> str:
    """Decode file bytes to string. Handles .docx and plain text."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return _read_docx(content)
    for enc in TEXT_ENCODINGS:
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {filename} as text (tried {TEXT_ENCODINGS})")


# -----------------------------------------------------------------------------
# RAG helpers
# -----------------------------------------------------------------------------
def _ensure_rag_services():
    global _vector_store, _doc_manager, _rag_import_error

    if _vector_store is not None and _doc_manager is not None:
        return _vector_store, _doc_manager

    if _rag_import_error is not None:
        raise RuntimeError(str(_rag_import_error))

    try:
        from rag.vector_store import VectorStore
        from rag.document_manager import DocumentManager
    except Exception as exc:
        _rag_import_error = exc
        raise RuntimeError(f"RAG dependencies unavailable: {exc}") from exc

    Path(VECTOR_DB_PATH).mkdir(parents=True, exist_ok=True)
    Path(UPLOADED_DOCS_PATH).mkdir(parents=True, exist_ok=True)
    _vector_store = VectorStore(VECTOR_DB_PATH)
    _doc_manager = DocumentManager(UPLOADED_DOCS_PATH)
    return _vector_store, _doc_manager


def _reset_rag_services_for_tests() -> None:
    global _vector_store, _doc_manager, _rag_import_error
    _vector_store = None
    _doc_manager = None
    _rag_import_error = None


# _openrouter_chat removed in favor of direct boto3 calls directly inside _run_bedrock_model


def _load_lmstudio_base_url() -> str:
    """Load LM Studio base URL from env, for connection only."""
    base_url = (_load_env_value("LMSTUDIO_BASE_URL") or LMSTUDIO_BASE_URL).strip().rstrip("/")
    if not base_url:
        base_url = LMSTUDIO_BASE_URL
    return base_url


def _load_local_model_config() -> tuple[str, str]:
    """Load the local inference connection config from env or local dotenv files."""
    base_url = _load_lmstudio_base_url()
    model = (_load_env_value("LOCAL_RAG_MODEL") or LOCAL_RAG_MODEL).strip()
    return base_url, (model or LOCAL_RAG_MODEL)


def _get_first_loaded_lmstudio_model() -> str:
    """Get the first loaded model from LM Studio, for legacy single-model endpoint."""
    base_url = _load_lmstudio_base_url()
    try:
        resp = requests.get(f"{base_url}/v1/models", timeout=2.0)
        if resp.status_code == 200:
            data = resp.json().get("data", [])
            if data:
                return data[0].get("id", "")
    except Exception:
        pass
    raise HTTPException(
        status_code=503,
        detail="No LM Studio model loaded. Please load a model in LM Studio first."
    )


def _load_opus_compare_model() -> str:
    model = (_load_env_value("OPUS_COMPARE_MODEL") or OPUS_COMPARE_MODEL).strip()
    return model or OPUS_COMPARE_MODEL


def _lmstudio_chat(messages: list[dict], model: str) -> dict:
    """Call LM Studio API with specified model."""
    if not model:
        raise ValueError("Model name is required for LM Studio calls")
    
    base_url = _load_lmstudio_base_url()
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.1,
        "stream": False,
    }

    try:
        resp = requests.post(f"{base_url}/v1/chat/completions", json=payload, timeout=240)
    except requests.RequestException as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Local LLM request failed: {exc}",
        ) from exc

    if resp.status_code != 200:
        try:
            err = resp.json()
            msg = err.get("error") or err.get("message") or resp.text
        except Exception:
            msg = resp.text
        raise HTTPException(status_code=resp.status_code, detail=msg)

    data = resp.json()
    content = str(data.get("choices", [{}])[0].get("message", {}).get("content", "")).strip()
    usage = data.get("usage", {}) if isinstance(data, dict) else {}
    return {
        "content": content,
        "usage": {
            "prompt_tokens": int(usage.get("prompt_tokens") or 0),
            "completion_tokens": int(usage.get("completion_tokens") or 0),
        },
        "model": data.get("model", model),
    }


def _build_metrics(
    prompt_tokens: int,
    completion_tokens: int,
    latency_s: float,
    eval_duration_s: Optional[float] = None,
) -> dict:
    total_tokens = int(prompt_tokens) + int(completion_tokens)
    tok_per_sec = 0.0
    if eval_duration_s and eval_duration_s > 0 and completion_tokens > 0:
        tok_per_sec = float(completion_tokens) / float(eval_duration_s)
    elif latency_s > 0 and completion_tokens > 0:
        tok_per_sec = float(completion_tokens) / float(latency_s)

    return {
        "total_tokens": total_tokens,
        "latency_ms": round(float(latency_s) * 1000.0, 1),
        "tok_per_sec": round(tok_per_sec, 2) if tok_per_sec > 0 else 0.0,
    }


def _build_rag_summary(collection_name: str, chunks: list[dict]) -> dict:
    docs_retrieved = {
        c.get("metadata", {}).get("doc_id")
        for c in chunks
        if c.get("metadata", {}).get("doc_id")
    }
    scores = [float(c.get("score", 0.0)) for c in chunks]
    return {
        "collection": collection_name,
        "chunks_retrieved": len(chunks),
        "documents_retrieved": len(docs_retrieved),
        "avg_score": round(sum(scores) / len(scores), 3) if scores else 0.0,
    }


def _build_sources(chunks: list[dict]) -> list[dict]:
    return [
        {
            "doc_id": c.get("metadata", {}).get("doc_id"),
            "chunk_id": c.get("metadata", {}).get("chunk_id"),
            "score": c.get("score", 0.0),
            "filename": c.get("metadata", {}).get("filename"),
            "snippet": (c.get("text", "") or "").strip()[:280],
        }
        for c in chunks
    ]


def _as_bool(value: Optional[str], default: bool = False) -> bool:
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _objective_uses_web_search(objective: Optional[str]) -> bool:
    return (objective or "").strip() in WEB_SEARCH_OBJECTIVES


def _build_recency_guard() -> str:
    today = datetime.now().date().isoformat()
    return (
        f"Current date: {today}. "
        "Do not present the output as if it were from 2025 by default. "
        "If you include a dated title such as 'Draft | YEAR', use the current year unless the user asks for a different year. "
        "For time-sensitive claims, use explicit month/year references."
    )


def _compact_text(value: Any) -> str:
    text = str(value or "")
    return " ".join(text.split()).strip()


def _query_exa_search(query: str) -> dict:
    key = _load_exa_key()
    if not key:
        return {
            "enabled": False,
            "provider": "exa",
            "error": "EXA_API_KEY not configured",
            "results": [],
        }

    headers = {
        "x-api-key": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
    }
    payload = {
        "query": _compact_text(query)[:500],
        "numResults": EXA_SEARCH_RESULT_LIMIT,
        "type": "fast",
    }

    started = time.perf_counter()
    try:
        resp = requests.post(
            EXA_SEARCH_URL,
            headers=headers,
            json=payload,
            timeout=EXA_SEARCH_TIMEOUT_S,
        )
    except requests.RequestException as exc:
        return {
            "enabled": False,
            "provider": "exa",
            "error": f"Exa request failed: {exc}",
            "results": [],
        }

    latency_ms = round((time.perf_counter() - started) * 1000.0, 1)
    if resp.status_code != 200:
        error_text = ""
        try:
            error_text = _compact_text(resp.json().get("error") or resp.json().get("message"))
        except Exception:
            error_text = _compact_text(resp.text)
        return {
            "enabled": False,
            "provider": "exa",
            "error": f"Exa returned {resp.status_code}: {error_text or 'unknown error'}",
            "results": [],
            "latency_ms": latency_ms,
        }

    try:
        data = resp.json()
    except Exception:
        return {
            "enabled": False,
            "provider": "exa",
            "error": "Exa returned non-JSON payload",
            "results": [],
            "latency_ms": latency_ms,
        }

    raw_results = data.get("results", []) if isinstance(data, dict) else []
    normalized = []
    for item in raw_results:
        if not isinstance(item, dict):
            continue
        url = _compact_text(item.get("url"))
        if not url:
            continue
        title = _compact_text(item.get("title")) or "Untitled result"
        published_date = _compact_text(
            item.get("publishedDate") or item.get("published_date") or item.get("date")
        )
        snippet = _compact_text(
            item.get("summary")
            or item.get("text")
            or item.get("snippet")
        )
        if len(snippet) > EXA_SNIPPET_CHARS:
            snippet = f"{snippet[:EXA_SNIPPET_CHARS].rstrip()}..."
        normalized.append(
            {
                "title": title,
                "url": url,
                "published_date": published_date,
                "snippet": snippet,
            }
        )

    return {
        "enabled": bool(normalized),
        "provider": "exa",
        "results": normalized,
        "latency_ms": latency_ms,
    }


def _build_web_context(
    query: str,
    objective: Optional[str],
    requested: bool,
) -> tuple[Optional[str], dict]:
    meta = {
        "requested": requested,
        "enabled": False,
        "provider": "exa",
        "results_count": 0,
        "sources": [],
    }

    if not requested:
        meta["reason"] = "disabled_by_user"
        return None, meta

    if not _objective_uses_web_search(objective):
        meta["reason"] = "objective_not_supported"
        return None, meta

    search = _query_exa_search(query=query)
    results = search.get("results") or []
    meta["enabled"] = bool(search.get("enabled"))
    meta["results_count"] = len(results)
    meta["sources"] = results
    if search.get("latency_ms") is not None:
        meta["latency_ms"] = search["latency_ms"]
    if search.get("error"):
        meta["error"] = search["error"]

    if not results:
        return None, meta

    rendered = []
    for idx, result in enumerate(results, start=1):
        published = result.get("published_date") or "Unknown"
        snippet = result.get("snippet") or "No snippet provided."
        rendered.append(
            f"[{idx}] {result.get('title', 'Untitled result')}\n"
            f"URL: {result.get('url', '')}\n"
            f"Published: {published}\n"
            f"Snippet: {snippet}"
        )
    web_context = "<web_context>\n" + "\n\n".join(rendered) + "\n</web_context>"
    return web_context, meta


def _run_bedrock_model(
    messages: list[dict],
    model_name: str,
    api_key: Optional[str] = None,
    enable_web_search: bool = False,
) -> dict:
    if not api_key:
        try:
            api_key = _load_bedrock_key()
        except RuntimeError as e:
            raise HTTPException(status_code=500, detail=str(e))
    # Provide the auth down to boto3 for API key users
    os.environ["AWS_BEARER_TOKEN_BEDROCK"] = api_key

    llm_started = time.perf_counter()
    system_prompt = ""
    filtered_messages = []
    for msg in messages:
        if msg["role"] == "system":
            system_prompt = msg["content"]
        else:
            filtered_messages.append(msg)

    try:
        client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
        body = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 4096,
            "messages": filtered_messages
        }
        if system_prompt:
            body["system"] = system_prompt

        response = client.invoke_model(
            modelId=model_name,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(body)
        )
        response_body = json.loads(response.get('body').read())
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    llm_elapsed_s = time.perf_counter() - llm_started
    content = response_body.get('content', [{}])[0].get('text', '').strip()
    usage = response_body.get('usage', {})
    prompt_tokens = int(usage.get("input_tokens", 0))
    completion_tokens = int(usage.get("output_tokens", 0))

    return {
        "content": content,
        "usage": {
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
        },
        "provider": "bedrock",
        "model": model_name,
        "metrics": _build_metrics(
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
            latency_s=llm_elapsed_s,
        ),
    }



def _run_lmstudio_model(messages: list[dict], model_name: str) -> dict:
    """Run LM Studio model with specified model name."""
    if not model_name:
        raise ValueError("Model name is required for LM Studio")
    
    # LM Studio models often don't support system role, so convert to user messages
    converted_messages = []
    system_content = ""
    
    for msg in messages:
        if msg["role"] == "system":
            system_content = msg["content"]
        else:
            converted_messages.append(msg)
    
    # Prepend system content to first user message if exists
    if system_content and converted_messages:
        for i, msg in enumerate(converted_messages):
            if msg["role"] == "user":
                converted_messages[i] = {
                    "role": "user",
                    "content": f"{system_content}\n\n{msg['content']}"
                }
                break
    
    llm_started = time.perf_counter()
    local = _lmstudio_chat(messages=converted_messages, model=model_name)
    llm_elapsed_s = time.perf_counter() - llm_started
    content = local.get("content", "")
    usage = local.get("usage", {})
    prompt_tokens = int(usage.get("prompt_tokens", 0))
    completion_tokens = int(usage.get("completion_tokens", 0))
    return {
        "content": content,
        "usage": {
            "prompt_tokens": prompt_tokens,
            "completion_tokens": completion_tokens,
        },
        "provider": "lmstudio",
        "model": local.get("model", model_name),
        "metrics": _build_metrics(
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
            latency_s=llm_elapsed_s,
        ),
    }


def _display_objective(value: Optional[str]) -> str:
    key = (value or "").strip()
    return OBJECTIVE_LABELS.get(key, key.replace("_", " ").title() or "Analysis")


def _display_mode(value: Optional[str]) -> str:
    key = (value or "").strip().lower()
    return MODE_LABELS.get(key, key.title() or "Standard")


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _format_timestamp(value: Any) -> str:
    raw = _safe_text(value)
    if raw:
        try:
            parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            return parsed.strftime("%B %d, %Y at %I:%M %p")
        except ValueError:
            return raw
    return datetime.now().strftime("%B %d, %Y at %I:%M %p")


def _slugify_filename(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug or "analysis"


def _build_export_filename(payload: dict[str, Any], extension: str) -> str:
    objective = _display_objective(payload.get("objective"))
    timestamp = datetime.now().strftime("%Y%m%d-%H%M")
    return f"benjamin-{_slugify_filename(objective)}-{timestamp}.{extension}"


def _parse_markdownish_blocks(text: str) -> list[dict[str, Any]]:
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    blocks: list[dict[str, Any]] = []
    idx = 0

    while idx < len(lines):
        current = lines[idx]
        stripped = current.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            if idx < len(lines):
                idx += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip()})
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": min(len(heading_match.group(1)), 4),
                    "text": heading_match.group(2).strip(),
                }
            )
            idx += 1
            continue

        if UL_PATTERN.match(current):
            items: list[str] = []
            while idx < len(lines):
                match = UL_PATTERN.match(lines[idx])
                if not match:
                    break
                items.append(match.group(1).strip())
                idx += 1
            blocks.append({"type": "ul", "items": items})
            continue

        if OL_PATTERN.match(current):
            items: list[str] = []
            while idx < len(lines):
                match = OL_PATTERN.match(lines[idx])
                if not match:
                    break
                items.append(match.group(1).strip())
                idx += 1
            blocks.append({"type": "ol", "items": items})
            continue

        paragraph_lines: list[str] = []
        while idx < len(lines):
            candidate = lines[idx]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if (
                candidate_stripped.startswith("```")
                or HEADING_PATTERN.match(candidate_stripped)
                or UL_PATTERN.match(candidate)
                or OL_PATTERN.match(candidate)
            ):
                break
            paragraph_lines.append(candidate.rstrip())
            idx += 1
        blocks.append({"type": "paragraph", "text": "\n".join(paragraph_lines).strip()})

    if not blocks and normalized.strip():
        return [{"type": "paragraph", "text": normalized.strip()}]
    return blocks


def _render_inline_html(text: str) -> str:
    raw = str(text or "")
    parts: list[str] = []
    cursor = 0
    for match in INLINE_TOKEN_PATTERN.finditer(raw):
        if match.start() > cursor:
            parts.append(escape(raw[cursor:match.start()]))

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            parts.append(f"<strong>{escape(token[2:-2])}</strong>")
        elif token.startswith("__") and token.endswith("__"):
            parts.append(f"<strong>{escape(token[2:-2])}</strong>")
        elif token.startswith("`") and token.endswith("`"):
            parts.append(
                "<code style=\"font-family: 'SFMono-Regular', Consolas, monospace; "
                "background:#eef3f8; border-radius:4px; padding:1px 4px;\">"
                f"{escape(token[1:-1])}</code>"
            )
        elif token.startswith("*") and token.endswith("*"):
            parts.append(f"<em>{escape(token[1:-1])}</em>")
        elif token.startswith("_") and token.endswith("_"):
            parts.append(f"<em>{escape(token[1:-1])}</em>")
        else:
            parts.append(escape(token))
        cursor = match.end()

    if cursor < len(raw):
        parts.append(escape(raw[cursor:]))

    return "".join(parts).replace("\n", "<br>")


def _render_blocks_html(text: str) -> str:
    blocks = _parse_markdownish_blocks(text)
    if not blocks:
        return "<p style=\"margin:0;\">No response text returned.</p>"

    html_parts: list[str] = []
    for block in blocks:
        block_type = block.get("type")
        if block_type == "heading":
            level = min(int(block.get("level", 2)) + 2, 6)
            html_parts.append(
                f"<h{level} style=\"margin:18px 0 8px; color:#10213c; font-size:{max(18, 28 - level * 2)}px;\">"
                f"{_render_inline_html(block.get('text', ''))}</h{level}>"
            )
        elif block_type == "ul":
            items = "".join(
                f"<li style=\"margin:0 0 6px;\">{_render_inline_html(item)}</li>"
                for item in block.get("items", [])
            )
            html_parts.append(f"<ul style=\"margin:10px 0 14px 22px; padding:0;\">{items}</ul>")
        elif block_type == "ol":
            items = "".join(
                f"<li style=\"margin:0 0 6px;\">{_render_inline_html(item)}</li>"
                for item in block.get("items", [])
            )
            html_parts.append(f"<ol style=\"margin:10px 0 14px 22px; padding:0;\">{items}</ol>")
        elif block_type == "code":
            html_parts.append(
                "<pre style=\"margin:10px 0 14px; padding:12px 14px; border-radius:10px; "
                "background:#0f172a; color:#e2e8f0; overflow:auto; font-size:13px; line-height:1.5;\">"
                f"{escape(block.get('text', ''))}</pre>"
            )
        else:
            html_parts.append(
                "<p style=\"margin:0 0 12px; color:#1f2f47; line-height:1.65;\">"
                f"{_render_inline_html(block.get('text', ''))}</p>"
            )
    return "".join(html_parts)


def _append_run_with_breaks(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    monospace: bool = False,
) -> None:
    parts = str(text or "").split("\n")
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.bold = bold
        run.italic = italic
        if monospace:
            run.font.name = "Courier New"
        if idx < len(parts) - 1:
            run.add_break()


def _append_formatted_runs(paragraph: Any, text: str) -> None:
    raw = str(text or "")
    cursor = 0
    for match in INLINE_TOKEN_PATTERN.finditer(raw):
        if match.start() > cursor:
            _append_run_with_breaks(paragraph, raw[cursor:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            _append_run_with_breaks(paragraph, token[2:-2], bold=True)
        elif token.startswith("__") and token.endswith("__"):
            _append_run_with_breaks(paragraph, token[2:-2], bold=True)
        elif token.startswith("`") and token.endswith("`"):
            _append_run_with_breaks(paragraph, token[1:-1], monospace=True)
        elif token.startswith("*") and token.endswith("*"):
            _append_run_with_breaks(paragraph, token[1:-1], italic=True)
        elif token.startswith("_") and token.endswith("_"):
            _append_run_with_breaks(paragraph, token[1:-1], italic=True)
        else:
            _append_run_with_breaks(paragraph, token)
        cursor = match.end()

    if cursor < len(raw):
        _append_run_with_breaks(paragraph, raw[cursor:])


def _append_markdownish_doc_content(document: Any, text: str) -> None:
    blocks = _parse_markdownish_blocks(text)
    if not blocks:
        document.add_paragraph("No response text returned.")
        return

    for block in blocks:
        block_type = block.get("type")
        if block_type == "heading":
            paragraph = document.add_heading(level=min(int(block.get("level", 2)) + 1, 4))
            _append_formatted_runs(paragraph, block.get("text", ""))
        elif block_type == "ul":
            for item in block.get("items", []):
                paragraph = document.add_paragraph(style="List Bullet")
                _append_formatted_runs(paragraph, item)
        elif block_type == "ol":
            for item in block.get("items", []):
                paragraph = document.add_paragraph(style="List Number")
                _append_formatted_runs(paragraph, item)
        elif block_type == "code":
            paragraph = document.add_paragraph()
            _append_run_with_breaks(paragraph, block.get("text", ""), monospace=True)
        else:
            paragraph = document.add_paragraph()
            _append_formatted_runs(paragraph, block.get("text", ""))


def _format_metric_value(label: str, value: Any) -> str:
    if value is None:
        return ""
    lower = label.lower()
    if lower.endswith("ms"):
        try:
            return f"{float(value):.1f} ms"
        except (TypeError, ValueError):
            return str(value)
    if lower.endswith("t/s"):
        try:
            return f"{float(value):.1f} t/s"
        except (TypeError, ValueError):
            return str(value)
    return str(value)


def _build_model_metric_rows(model_payload: dict[str, Any]) -> list[tuple[str, str]]:
    usage = model_payload.get("usage") if isinstance(model_payload, dict) else {}
    metrics = model_payload.get("metrics") if isinstance(model_payload, dict) else {}
    rows: list[tuple[str, str]] = []

    provider = _safe_text(model_payload.get("provider"))
    model_name = _safe_text(model_payload.get("model"))
    if provider:
        rows.append(("Provider", provider))
    if model_name:
        rows.append(("Model", model_name))

    prompt_tokens = usage.get("prompt_tokens") if isinstance(usage, dict) else None
    completion_tokens = usage.get("completion_tokens") if isinstance(usage, dict) else None
    total_tokens = metrics.get("total_tokens") if isinstance(metrics, dict) else None
    latency_ms = metrics.get("latency_ms") if isinstance(metrics, dict) else None
    tok_per_sec = metrics.get("tok_per_sec") if isinstance(metrics, dict) else None
    retrieval_ms = metrics.get("retrieval_ms") if isinstance(metrics, dict) else None

    if prompt_tokens not in (None, ""):
        rows.append(("Prompt tokens", str(prompt_tokens)))
    if completion_tokens not in (None, ""):
        rows.append(("Completion tokens", str(completion_tokens)))
    if total_tokens not in (None, ""):
        rows.append(("Total tokens", str(total_tokens)))
    if latency_ms not in (None, "", 0):
        rows.append(("Latency ms", _format_metric_value("latency ms", latency_ms)))
    if tok_per_sec not in (None, "", 0):
        rows.append(("Speed t/s", _format_metric_value("speed t/s", tok_per_sec)))
    if retrieval_ms not in (None, "", 0):
        rows.append(("Retrieval ms", _format_metric_value("retrieval ms", retrieval_ms)))

    return rows


def _render_key_value_table_html(rows: list[tuple[str, str]]) -> str:
    filtered_rows = [(label, value) for label, value in rows if value]
    if not filtered_rows:
        return ""

    table_rows = "".join(
        "<tr>"
        f"<th align=\"left\" style=\"padding:8px 10px; border:1px solid #d6deea; background:#eef3f8; width:180px;\">{escape(label)}</th>"
        f"<td style=\"padding:8px 10px; border:1px solid #d6deea;\">{escape(value)}</td>"
        "</tr>"
        for label, value in filtered_rows
    )
    return (
        "<table role=\"presentation\" cellspacing=\"0\" cellpadding=\"0\" "
        "style=\"border-collapse:collapse; width:100%; margin:10px 0 16px; font-size:14px;\">"
        f"{table_rows}</table>"
    )


def _coerce_context_fields(payload: dict[str, Any]) -> list[tuple[str, str]]:
    raw_fields = payload.get("context_fields")
    if not isinstance(raw_fields, list):
        return []

    fields: list[tuple[str, str]] = []
    for item in raw_fields:
        if not isinstance(item, dict):
            continue
        label = _safe_text(item.get("label"))
        value = _safe_text(item.get("value"))
        if label and value:
            fields.append((label, value))
    return fields


def _coerce_sources(payload: dict[str, Any]) -> list[dict[str, Any]]:
    raw_sources = payload.get("sources")
    if not isinstance(raw_sources, list):
        return []
    return [source for source in raw_sources if isinstance(source, dict)]


def _coerce_web_sources(payload: dict[str, Any]) -> list[dict[str, Any]]:
    web_meta = payload.get("web_search")
    if not isinstance(web_meta, dict):
        return []
    raw_sources = web_meta.get("sources")
    if not isinstance(raw_sources, list):
        return []
    return [source for source in raw_sources if isinstance(source, dict)]


def _build_export_html(payload: dict[str, Any]) -> str:
    objective_label = _display_objective(payload.get("objective"))
    mode_label = _display_mode(payload.get("mode"))
    generated_at = _format_timestamp(payload.get("generated_at"))
    prompt = _safe_text(payload.get("prompt")) or "No prompt captured."
    left = payload.get("left") if isinstance(payload.get("left"), dict) else {}
    right = payload.get("right") if isinstance(payload.get("right"), dict) else {}

    context_rows = [
        ("Module", objective_label),
        ("Context mode", mode_label),
        ("Exported", generated_at),
    ]
    context_rows.extend(_coerce_context_fields(payload))

    rag = payload.get("rag") if isinstance(payload.get("rag"), dict) else {}
    rag_rows: list[tuple[str, str]] = []
    if rag:
        rag_rows = [
            ("Collection", _safe_text(rag.get("collection")) or "n/a"),
            ("Chunks retrieved", str(rag.get("chunks_retrieved", 0))),
            ("Documents retrieved", str(rag.get("documents_retrieved", 0))),
        ]
        try:
            rag_rows.append(("Average confidence", f"{round(float(rag.get('avg_score', 0.0)) * 100)}%"))
        except (TypeError, ValueError):
            pass

    left_title = _safe_text(left.get("export_label")) or "Left Model"
    right_title = _safe_text(right.get("export_label")) or "Right Model"
    sources = _coerce_sources(payload)
    web_sources = _coerce_web_sources(payload)

    source_html = ""
    if sources:
        source_items = []
        for source in sources:
            confidence_label = "n/a"
            try:
                confidence_label = f"{round(float(source.get('score', 0.0)) * 100)}%"
            except (TypeError, ValueError):
                pass
            chunk_id = source.get("chunk_id")
            chunk_label = f"Fragment {chunk_id + 1}" if isinstance(chunk_id, int) else "Fragment"
            snippet = _safe_text(source.get("snippet"))
            source_items.append(
                "<li style=\"margin:0 0 12px; padding:12px 14px; border:1px solid #d6deea; border-radius:10px; background:#ffffff;\">"
                f"<div style=\"font-weight:700; color:#10213c;\">{escape(_safe_text(source.get('filename')) or 'Untitled source')}</div>"
                f"<div style=\"margin-top:4px; color:#4d5e78; font-size:13px;\">{escape(chunk_label)} | Confidence {escape(confidence_label)}</div>"
                f"{f'<div style=\"margin-top:8px; color:#1f2f47;\">{escape(snippet)}</div>' if snippet else ''}"
                "</li>"
            )
        source_html = (
            "<section style=\"margin-top:28px;\">"
            "<h2 style=\"margin:0 0 12px; color:#10213c; font-size:22px;\">Referenced Passages</h2>"
            f"<ul style=\"list-style:none; margin:0; padding:0;\">{''.join(source_items)}</ul>"
            "</section>"
        )

    web_html = ""
    if web_sources:
        web_items = []
        for source in web_sources:
            title = _safe_text(source.get("title")) or "Untitled source"
            url = _safe_text(source.get("url"))
            published = _safe_text(source.get("published_date")) or "Unknown date"
            snippet = _safe_text(source.get("snippet"))
            link_html = (
                f"<div style=\"margin-top:6px;\"><a href=\"{escape(url)}\" style=\"color:#1d4ed8; text-decoration:none;\">{escape(url)}</a></div>"
                if url
                else ""
            )
            web_items.append(
                "<li style=\"margin:0 0 12px; padding:12px 14px; border:1px solid #d6deea; border-radius:10px; background:#ffffff;\">"
                f"<div style=\"font-weight:700; color:#10213c;\">{escape(title)}</div>"
                f"<div style=\"margin-top:4px; color:#4d5e78; font-size:13px;\">{escape(published)}</div>"
                f"{link_html}"
                f"{f'<div style=\"margin-top:8px; color:#1f2f47;\">{escape(snippet)}</div>' if snippet else ''}"
                "</li>"
            )
        web_html = (
            "<section style=\"margin-top:28px;\">"
            "<h2 style=\"margin:0 0 12px; color:#10213c; font-size:22px;\">Live Web Context</h2>"
            f"<ul style=\"list-style:none; margin:0; padding:0;\">{''.join(web_items)}</ul>"
            "</section>"
        )

    rag_html = (
        f"<section style=\"margin-top:28px;\"><h2 style=\"margin:0 0 12px; color:#10213c; font-size:22px;\">RAG Summary</h2>{_render_key_value_table_html(rag_rows)}</section>"
        if rag_rows
        else ""
    )

    return f"""<!DOCTYPE html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Benjamin Maurice Analysis Export</title>
  </head>
  <body style="margin:0; padding:24px; background:#eef2f7; color:#132238; font-family:Calibri, Arial, sans-serif;">
    <div style="max-width:960px; margin:0 auto; background:#ffffff; border:1px solid #d6deea; border-radius:18px; overflow:hidden;">
      <div style="padding:28px 32px; background:linear-gradient(135deg, #10213c, #274b7a); color:#ffffff;">
        <div style="font-size:12px; letter-spacing:0.12em; text-transform:uppercase; opacity:0.82;">Benjamin Maurice</div>
        <h1 style="margin:8px 0 0; font-size:32px; line-height:1.15;">Analysis Export</h1>
        <div style="margin-top:10px; font-size:15px; opacity:0.9;">{escape(objective_label)} | {escape(generated_at)}</div>
      </div>
      <div style="padding:28px 32px 34px;">
        <section>
          <h2 style="margin:0 0 12px; color:#10213c; font-size:22px;">Run Summary</h2>
          {_render_key_value_table_html(context_rows)}
        </section>

        <section style="margin-top:24px;">
          <h2 style="margin:0 0 12px; color:#10213c; font-size:22px;">Prompt</h2>
          <div style="padding:16px 18px; border:1px solid #d6deea; border-radius:12px; background:#f8fafc; color:#1f2f47; line-height:1.7;">
            {_render_blocks_html(prompt)}
          </div>
        </section>

        <section style="margin-top:28px;">
          <h2 style="margin:0 0 14px; color:#10213c; font-size:22px;">Model Outputs</h2>

          <div style="margin:0 0 18px; padding:18px; border:1px solid #d6deea; border-radius:14px; background:#fbfdff;">
            <h3 style="margin:0 0 8px; color:#10213c; font-size:20px;">{escape(left_title)}</h3>
            {_render_key_value_table_html(_build_model_metric_rows(left))}
            {_render_blocks_html(_safe_text(left.get("content")) or "No response text returned.")}
          </div>

          <div style="margin:0; padding:18px; border:1px solid #d6deea; border-radius:14px; background:#fbfdff;">
            <h3 style="margin:0 0 8px; color:#10213c; font-size:20px;">{escape(right_title)}</h3>
            {_render_key_value_table_html(_build_model_metric_rows(right))}
            {_render_blocks_html(_safe_text(right.get("content")) or "No response text returned.")}
          </div>
        </section>

        {rag_html}
        {source_html}
        {web_html}
      </div>
    </div>
  </body>
</html>"""


def _build_export_docx(payload: dict[str, Any]) -> bytes:
    from docx import Document

    objective_label = _display_objective(payload.get("objective"))
    mode_label = _display_mode(payload.get("mode"))
    generated_at = _format_timestamp(payload.get("generated_at"))
    prompt = _safe_text(payload.get("prompt")) or "No prompt captured."
    left = payload.get("left") if isinstance(payload.get("left"), dict) else {}
    right = payload.get("right") if isinstance(payload.get("right"), dict) else {}

    document = Document()
    document.core_properties.title = "Benjamin Maurice Analysis Export"
    document.core_properties.subject = objective_label

    title = document.add_heading("Benjamin Maurice Analysis Export", level=0)
    if title.runs:
        title.runs[0].bold = True
    document.add_paragraph(f"{objective_label} | Exported {generated_at}")

    document.add_heading("Run Summary", level=1)
    summary_rows = [
        ("Module", objective_label),
        ("Context mode", mode_label),
        ("Exported", generated_at),
        *_coerce_context_fields(payload),
    ]
    for label, value in summary_rows:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_heading("Prompt", level=1)
    _append_markdownish_doc_content(document, prompt)

    for section_title, model_payload in (
        (_safe_text(left.get("export_label")) or "Left Model", left),
        (_safe_text(right.get("export_label")) or "Right Model", right),
    ):
        document.add_heading(section_title, level=1)
        for label, value in _build_model_metric_rows(model_payload):
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)
        _append_markdownish_doc_content(
            document,
            _safe_text(model_payload.get("content")) or "No response text returned.",
        )

    rag = payload.get("rag") if isinstance(payload.get("rag"), dict) else {}
    if rag:
        document.add_heading("RAG Summary", level=1)
        rag_rows = [
            ("Collection", _safe_text(rag.get("collection")) or "n/a"),
            ("Chunks retrieved", str(rag.get("chunks_retrieved", 0))),
            ("Documents retrieved", str(rag.get("documents_retrieved", 0))),
        ]
        try:
            rag_rows.append(("Average confidence", f"{round(float(rag.get('avg_score', 0.0)) * 100)}%"))
        except (TypeError, ValueError):
            pass
        for label, value in rag_rows:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)

    sources = _coerce_sources(payload)
    if sources:
        document.add_heading("Referenced Passages", level=1)
        for source in sources:
            heading = document.add_paragraph()
            heading.add_run(_safe_text(source.get("filename")) or "Untitled source").bold = True
            chunk_id = source.get("chunk_id")
            if isinstance(chunk_id, int):
                heading.add_run(f" | Fragment {chunk_id + 1}")
            try:
                heading.add_run(f" | Confidence {round(float(source.get('score', 0.0)) * 100)}%")
            except (TypeError, ValueError):
                pass
            snippet = _safe_text(source.get("snippet"))
            if snippet:
                document.add_paragraph(snippet)

    web_sources = _coerce_web_sources(payload)
    if web_sources:
        document.add_heading("Live Web Context", level=1)
        for source in web_sources:
            paragraph = document.add_paragraph()
            paragraph.add_run(_safe_text(source.get("title")) or "Untitled source").bold = True
            published = _safe_text(source.get("published_date"))
            if published:
                paragraph.add_run(f" | {published}")
            url = _safe_text(source.get("url"))
            if url:
                document.add_paragraph(url)
            snippet = _safe_text(source.get("snippet"))
            if snippet:
                document.add_paragraph(snippet)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# -----------------------------------------------------------------------------
# Routes
# -----------------------------------------------------------------------------
@app.post("/api/chat")
async def chat(
    message: str = Form(..., description="User message/prompt"),
    file: Optional[UploadFile] = File(None),
    mode: str = Form("direct"),
    objective: Optional[str] = Form(None),
    top_k: int = Form(5),
    min_score: float = Form(0.5),
    use_web_search: Optional[str] = Form("false"),
):
    """
    Chat endpoint with two modes:
    - direct: existing file-upload behavior
    - rag: retrieve from local vector store and include sources
    """
    mode = (mode or "direct").strip().lower()
    request_started = time.perf_counter()
    web_requested = _as_bool(use_web_search, default=False)
    web_context, web_meta = _build_web_context(
        query=message.strip(),
        objective=objective,
        requested=web_requested,
    )
    if mode not in {"direct", "rag"}:
        raise HTTPException(status_code=400, detail="mode must be 'direct' or 'rag'")

    if mode == "direct":
        # Build user content: optional file context + message
        user_content_parts = []

        if file and file.filename:
            raw = await file.read()
            if len(raw) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail=f"File too large (max {MAX_FILE_SIZE // 1024}KB)",
                )
            try:
                text = _read_file_content(raw, file.filename)
            except ValueError as e:
                raise HTTPException(status_code=400, detail=str(e))
            user_content_parts.append(f"<file name=\"{file.filename}\">\n{text}\n</file>")

        if web_context:
            user_content_parts.append(web_context)
        user_content_parts.append(message.strip())
        user_content = "\n\n" + "\n\n".join(user_content_parts)
        direct_messages = [{"role": "user", "content": user_content}]
        if _objective_uses_web_search(objective):
            direct_messages = [
                {"role": "system", "content": _build_recency_guard()},
                {"role": "user", "content": user_content},
            ]
        response = _run_bedrock_model(
            messages=direct_messages,
            model_name=MODEL,
            enable_web_search=web_meta["enabled"],
        )
        response["web_search"] = web_meta
        return response

    if not objective or objective not in OBJECTIVE_TO_COLLECTION:
        raise HTTPException(
            status_code=400,
            detail="objective must be 'expert_network_brief', 'interview_guide', or 'insights_qa' in rag mode",
        )

    try:
        vector_store, _ = _ensure_rag_services()
        from rag.retrieval import assemble_context, retrieve_context
        from rag.system_prompts import load_system_prompt
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    collection_name = OBJECTIVE_TO_COLLECTION[objective]
    retrieval_started = time.perf_counter()
    chunks = retrieve_context(
        query=message.strip(),
        vector_store=vector_store,
        collection_name=collection_name,
        top_k=max(1, min(top_k, 20)),
        min_score=min_score,
    )
    retrieval_elapsed_s = time.perf_counter() - retrieval_started
    if not chunks:
        raise HTTPException(
            status_code=400,
            detail=f"No relevant context found in collection '{collection_name}'",
        )

    context = assemble_context(chunks)
    system_prompt = load_system_prompt(objective)
    if _objective_uses_web_search(objective):
        system_prompt = f"{system_prompt}\n\n{_build_recency_guard()}"
    user_sections = [
        "<retrieved_context>\n"
        f"{context}\n"
        "</retrieved_context>"
    ]
    if web_context:
        user_sections.append(web_context)
    user_sections.append(f"User request:\n{message.strip()}")
    user_content = "\n\n".join(user_sections)

    provider = OBJECTIVE_TO_PROVIDER.get(objective, "bedrock")
    model_response = None
    if provider == "bedrock":
        model_response = _run_bedrock_model(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            model_name=MODEL,
            enable_web_search=web_meta["enabled"],
        )
    elif provider == "lmstudio":
        loaded_model = _get_first_loaded_lmstudio_model()
        model_response = _run_lmstudio_model(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            model_name=loaded_model,
        )
    else:
        raise HTTPException(status_code=500, detail=f"Unsupported provider: {provider}")

    rag_summary = _build_rag_summary(collection_name=collection_name, chunks=chunks)
    sources = _build_sources(chunks)
    request_elapsed_s = time.perf_counter() - request_started
    model_response["metrics"]["retrieval_ms"] = round(retrieval_elapsed_s * 1000.0, 1)
    model_response["metrics"]["request_total_ms"] = round(request_elapsed_s * 1000.0, 1)

    return {
        **model_response,
        "rag": rag_summary,
        "sources": sources,
        "web_search": web_meta,
    }

def _get_available_models():
    models_path = Path(APP_DIR) / "models.txt"
    models = []
    seen = set()
    
    if models_path.exists():
        for line in models_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or "|" not in line:
                continue
            provider, model = line.split("|", 1)
            provider, model = provider.strip(), model.strip()
            models.append({"provider": provider, "model": model})
            seen.add(f"{provider}|{model}")

    # Dynamically fetch LM Studio models
    base_url = _load_lmstudio_base_url()
    
    try:
        resp = requests.get(f"{base_url}/v1/models", timeout=2.0)
        if resp.status_code == 200:
            data = resp.json().get("data", [])
            for m in data:
                model_name = m.get("id")
                if model_name:
                    key = f"lmstudio|{model_name}"
                    if key not in seen:
                        models.append({"provider": "lmstudio", "model": model_name})
                        seen.add(key)
    except Exception:
        pass  # LM Studio might not be running or reachable, which is fine

    return models

@app.get("/api/models")
async def get_models():
    """Returns the list of available models from models.txt."""
    return {"models": _get_available_models()}


@app.post("/api/chat/compare")
async def chat_compare(
    message: str = Form(..., description="User message/prompt"),
    mode: Optional[str] = Form("direct"),
    objective: str = Form(..., description="Objective to run the prompt for"),
    model_left: str = Form(..., description="Provider|Model for the left pane"),
    model_right: str = Form(..., description="Provider|Model for the right pane"),
    top_k: int = Form(5),
    min_score: float = Form(0.5),
    file: Optional[UploadFile] = File(None),
    use_web_search: Optional[str] = Form("false"),
):
    """
    Compare endpoint for side-by-side answers for any objective.
    Uses model_left and model_right format: "provider|model".
    """
    mode = (mode or "direct").strip().lower()
    request_started = time.perf_counter()
    objective = (objective or "").strip()
    
    rag_summary = None
    sources = []
    retrieval_elapsed_s = 0
    web_requested = _as_bool(use_web_search, default=False)
    web_context, web_meta = _build_web_context(
        query=message.strip(),
        objective=objective,
        requested=web_requested,
    )

    if mode == "direct":
        user_content_parts = []
        if file and file.filename:
            raw = await file.read()
            if len(raw) > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE // 1024}KB)")
            try:
                text = _read_file_content(raw, file.filename)
            except ValueError as e:
                raise HTTPException(status_code=400, detail=str(e))
            user_content_parts.append(f"<file name=\"{file.filename}\">\n{text}\n</file>")
        if web_context:
            user_content_parts.append(web_context)
        user_content_parts.append(message.strip())
        user_content = "\n\n" + "\n\n".join(user_content_parts)
        shared_messages = [{"role": "user", "content": user_content}]
        if _objective_uses_web_search(objective):
            shared_messages = [
                {"role": "system", "content": _build_recency_guard()},
                {"role": "user", "content": user_content},
            ]
    else:
        if objective not in OBJECTIVE_TO_COLLECTION:
            raise HTTPException(status_code=400, detail="invalid objective")

        try:
            vector_store, _ = _ensure_rag_services()
            from rag.retrieval import assemble_context, retrieve_context
            from rag.system_prompts import load_system_prompt
        except RuntimeError as exc:
            raise HTTPException(status_code=500, detail=str(exc))

        collection_name = OBJECTIVE_TO_COLLECTION[objective]
        retrieval_started = time.perf_counter()
        chunks = retrieve_context(
            query=message.strip(),
            vector_store=vector_store,
            collection_name=collection_name,
            top_k=max(1, min(top_k, 20)),
            min_score=min_score,
        )
        retrieval_elapsed_s = time.perf_counter() - retrieval_started
        if not chunks:
            raise HTTPException(status_code=400, detail=f"No relevant context found in collection '{collection_name}'")

        context = assemble_context(chunks)
        system_prompt = load_system_prompt(objective)
        if _objective_uses_web_search(objective):
            system_prompt = f"{system_prompt}\n\n{_build_recency_guard()}"
        user_sections = [
            "<retrieved_context>\n"
            f"{context}\n"
            "</retrieved_context>"
        ]
        if web_context:
            user_sections.append(web_context)
        user_sections.append(f"User request:\n{message.strip()}")
        user_content = "\n\n".join(user_sections)
        shared_messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ]
        rag_summary = _build_rag_summary(collection_name=collection_name, chunks=chunks)
        sources = _build_sources(chunks)

    def _run_selected_model(model_identifier, msgs):
        try:
            prov, mod = model_identifier.split("|", 1)
        except ValueError:
            return {"error": "Invalid model identifier format"}

        if prov == "bedrock":
            try:
                bedrock_key = _load_bedrock_key()
                return _run_bedrock_model(
                    msgs,
                    mod,
                    bedrock_key,
                    enable_web_search=web_meta["enabled"],
                )
            except Exception as e:
                return {"error": str(e), "content": f"Error: {e}"}
        elif prov == "lmstudio":
            try:
                return _run_lmstudio_model(msgs, model_name=mod)
            except Exception as e:
                return {"error": str(e), "content": f"Error: {e}"}
        else:
            return {"error": "Unknown provider", "content": f"Unknown provider {prov}"}

    with ThreadPoolExecutor(max_workers=2) as executor:
        local_future = executor.submit(_run_selected_model, model_left, shared_messages)
        opus_future = executor.submit(_run_selected_model, model_right, shared_messages)
        
        local_result = local_future.result()
        opus_result = opus_future.result()

    retrieval_ms = round(retrieval_elapsed_s * 1000.0, 1) if retrieval_elapsed_s else 0
    request_elapsed_s = time.perf_counter() - request_started
    
    if "metrics" in local_result and retrieval_ms:
        local_result["metrics"]["retrieval_ms"] = retrieval_ms
    if "metrics" in opus_result and retrieval_ms:
        opus_result["metrics"]["retrieval_ms"] = retrieval_ms

    return {
        "left": local_result,
        "right": opus_result,
        "rag": rag_summary,
        "sources": sources,
        "web_search": web_meta,
        "metrics": {
            "retrieval_ms": retrieval_ms,
            "request_total_ms": round(request_elapsed_s * 1000.0, 1),
        },
    }


@app.post("/api/export")
async def export_analysis(payload: dict[str, Any]):
    export_format = _safe_text(payload.get("export_format")).lower()
    if export_format not in {"word", "email_html"}:
        raise HTTPException(
            status_code=400,
            detail="export_format must be 'word' or 'email_html'",
        )

    left = payload.get("left")
    right = payload.get("right")
    if not isinstance(left, dict) or not isinstance(right, dict):
        raise HTTPException(
            status_code=400,
            detail="left and right model payloads are required for export",
        )

    if export_format == "email_html":
        content = _build_export_html(payload)
        filename = _build_export_filename(payload, "html")
        return Response(
            content=content,
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    content = _build_export_docx(payload)
    filename = _build_export_filename(payload, "docx")
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/documents")
def documents(vertical: Optional[str] = None):
    try:
        _, doc_manager = _ensure_rag_services()
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))
    return doc_manager.list_documents(vertical=vertical)


@app.post("/api/documents/sync")
def sync_documents():
    """Sync the RAG backend with the files currently in the Data/ directory."""
    try:
        vector_store, doc_manager = _ensure_rag_services()
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))
        
    data_dir = os.path.join(APP_DIR, "Data")
    if not os.path.exists(data_dir):
        return {"ingested": 0, "message": "Data directory not found."}
        
    try:
        files = _resolve_files(None, data_dir)
        ingested, removed = sync_collection(UNIFIED_COLLECTION, files, vector_store, doc_manager)
        return {"ingested": ingested, "removed": removed, "message": "Sync successful."}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Sync failed: {str(exc)}")


@app.get("/api/documents/{doc_id}/file")
def document_file(doc_id: str):
    try:
        _, doc_manager = _ensure_rag_services()
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    doc = doc_manager.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"Document '{doc_id}' not found")

    stored_path = Path(str(doc.get("stored_path", "")))
    if not stored_path.exists() or not stored_path.is_file():
        raise HTTPException(
            status_code=404,
            detail=f"Stored file missing for document '{doc_id}'",
        )

    return FileResponse(
        path=str(stored_path),
        filename=doc.get("filename") or stored_path.name,
    )


@app.get("/api/stats")
def stats():
    try:
        vector_store, doc_manager = _ensure_rag_services()
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    docs = doc_manager.list_documents()
    total_chunks = sum(int(d["chunk_count"]) for d in docs)
    by_vertical = {}
    for d in docs:
        vertical = d["vertical"]
        entry = by_vertical.setdefault(vertical, {"doc_count": 0, "chunk_count": 0})
        entry["doc_count"] += 1
        entry["chunk_count"] += int(d["chunk_count"])

    for vertical in by_vertical:
        by_vertical[vertical]["vector_count"] = vector_store.get_stats(vertical)["vector_count"]

    return {
        "total_documents": len(docs),
        "total_chunks": total_chunks,
        "by_vertical": by_vertical,
    }


@app.get("/")
def index():
    """Serve the legacy classic frontend."""
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Classic theme not found."}


@app.get("/classic")
def index_classic():
    """Serve the legacy classic frontend."""
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Classic theme not found."}


@app.get("/ultra")
def index_ultra():
    """Serve the ultra premium frontend."""
    ultra_path = os.path.join(STATIC_DIR, "index_ultra.html")
    if os.path.exists(ultra_path):
        return FileResponse(ultra_path)
    return {"message": "Ultra theme not found."}


# Mount static files (for CSS, JS if any)
if os.path.exists(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
